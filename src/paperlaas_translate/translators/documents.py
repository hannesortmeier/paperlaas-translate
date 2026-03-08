from __future__ import annotations

from dataclasses import dataclass
import shlex
import subprocess
import re
from collections.abc import Iterable, Sequence
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from uuid import uuid4

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..llm import OpenAISegmentTranslator, SegmentTranslationError
from ..logging_config import get_logger
from ..models import TargetLanguage
from ..utils import UnsupportedDocumentTypeError, is_docx, is_legacy_word, is_odt, is_text

logger = get_logger(__name__)

_PARAGRAPH_BREAK_RE = re.compile(r"(\n\s*\n)")
_SEGMENT_MARKER_TOKEN_RE = re.compile(r"\[\[PAPERLAAS_SEGMENT_BREAK_[^\]]+\]\]")
_LEADING_TRAILING_WHITESPACE_RE = re.compile(r"^([ \t]*)(.*?)([ \t]*)$", re.DOTALL)
_TAB_TOKEN = "[[PAPERLAAS_TAB]]"
_LINE_BREAK_TOKEN = "[[PAPERLAAS_LINE_BREAK]]"
_FORMATTING_RETRY_LIMIT = 2


@dataclass(frozen=True, slots=True)
class _PreparedSegment:
    leading_whitespace: str
    encoded_body: str
    trailing_whitespace: str
    tab_count: int
    line_break_count: int


class _FormattingPreservationError(SegmentTranslationError):
    pass


class OfficeAndTextTranslator:
    def __init__(
        self,
        segment_translator: OpenAISegmentTranslator,
        *,
        batch_chars: int,
        batch_items: int,
        converter_command: str = "soffice",
        temp_dir: str = "/tmp/paperlaas-translate",
        conversion_timeout_seconds: float = 120.0,
    ) -> None:
        self._segment_translator = segment_translator
        self._batch_chars = batch_chars
        self._batch_items = batch_items
        self._converter_command = converter_command
        self._temp_dir = Path(temp_dir)
        self._conversion_timeout_seconds = conversion_timeout_seconds

    def translate(self, mime_type: str, content: bytes, target_language: TargetLanguage) -> bytes:
        if is_docx(mime_type):
            return self._translate_docx(content, target_language)
        if is_odt(mime_type):
            return self._translate_odt(content, target_language)
        if is_text(mime_type):
            return self._translate_text(content, target_language)
        if is_legacy_word(mime_type):
            raise UnsupportedDocumentTypeError(
                "Legacy Word .doc files require an external converter and are not supported yet"
            )
        raise UnsupportedDocumentTypeError(f"Unsupported office/text mime type: {mime_type}")

    def _translate_docx(self, content: bytes, target_language: TargetLanguage) -> bytes:
        document = Document(BytesIO(content))
        paragraphs = [paragraph for paragraph in _iter_document_paragraphs(document) if paragraph.text.strip()]
        translated_texts = self._translate_batches(
            [paragraph.text for paragraph in paragraphs],
            target_language,
        )

        for paragraph, translated_text in zip(paragraphs, translated_texts, strict=True):
            _replace_paragraph_text(paragraph, translated_text)

        output = BytesIO()
        document.save(output)
        logger.info(
            "translated docx document",
            paragraph_count=len(paragraphs),
            target_language=target_language.value,
        )
        return output.getvalue()

    def _translate_odt(self, content: bytes, target_language: TargetLanguage) -> bytes:
        self._temp_dir.mkdir(parents=True, exist_ok=True)
        with TemporaryDirectory(prefix="office-", dir=self._temp_dir) as temp_dir_name:
            temp_dir = Path(temp_dir_name)
            source_odt_path = temp_dir / "source.odt"
            source_odt_path.write_bytes(content)

            source_docx_path = self._convert_office_document(
                source_odt_path,
                "docx",
                temp_dir,
            )
            translated_docx_bytes = self._translate_docx(
                source_docx_path.read_bytes(),
                target_language,
            )

            translated_docx_path = temp_dir / "translated.docx"
            translated_docx_path.write_bytes(translated_docx_bytes)
            translated_odt_path = self._convert_office_document(
                translated_docx_path,
                "odt",
                temp_dir,
            )
            logger.info(
                "translated odt document",
                target_language=target_language.value,
                output_file=str(translated_odt_path),
            )
            return translated_odt_path.read_bytes()

    def _translate_text(self, content: bytes, target_language: TargetLanguage) -> bytes:
        decoded = _decode_text(content)
        parts = _PARAGRAPH_BREAK_RE.split(decoded)
        translatable_blocks = [part for index, part in enumerate(parts) if index % 2 == 0 and part.strip()]
        translated_blocks = iter(self._translate_batches(translatable_blocks, target_language))

        rebuilt: list[str] = []
        for index, part in enumerate(parts):
            if index % 2 == 0 and part.strip():
                rebuilt.append(next(translated_blocks))
            else:
                rebuilt.append(part)

        translated = "".join(rebuilt)
        logger.info(
            "translated text document",
            source_length=len(decoded),
            output_length=len(translated),
            target_language=target_language.value,
        )
        return translated.encode("utf-8")

    def _translate_batches(
        self,
        segments: Sequence[str],
        target_language: TargetLanguage,
    ) -> list[str]:
        if not segments:
            return []

        translated_output: list[str] = []
        current_chunk: list[str] = []
        current_chars = 0
        max_chunk_chars = max(self._batch_chars, self._batch_chars * self._batch_items)

        for segment in segments:
            segment_length = len(segment)
            if current_chunk and current_chars + segment_length > max_chunk_chars:
                translated_output.extend(
                    self._translate_chunk_with_fallback(current_chunk, target_language)
                )
                current_chunk = []
                current_chars = 0
            current_chunk.append(segment)
            current_chars += segment_length

        if current_chunk:
            translated_output.extend(self._translate_chunk_with_fallback(current_chunk, target_language))

        return translated_output

    def _translate_chunk_with_fallback(
        self,
        segments: Sequence[str],
        target_language: TargetLanguage,
    ) -> list[str]:
        marker = _make_segment_marker()
        prepared_segments = [_prepare_segment_for_translation(segment) for segment in segments]
        payload = marker.join(segment.encoded_body for segment in prepared_segments)
        chunk_stats = _get_chunk_stats(segments)
        try:
            formatting_retry_count = 0
            while True:
                translated_payload = self._segment_translator.translate_segments([payload], target_language)[0]
                translated_segments = _split_translated_payload(translated_payload, marker, len(segments))
                try:
                    return [
                        _restore_translated_segment(prepared_segment, translated_segment)
                        for prepared_segment, translated_segment in zip(
                            prepared_segments, translated_segments, strict=True
                        )
                    ]
                except _FormattingPreservationError as exc:
                    failure_context = {
                        "chunk_size": len(segments),
                        "payload_length": len(payload),
                        "failure_reason": str(exc),
                        "retry_count": formatting_retry_count,
                        "retry_limit": _FORMATTING_RETRY_LIMIT,
                        "target_language": target_language.value,
                        **chunk_stats,
                    }
                    if formatting_retry_count < _FORMATTING_RETRY_LIMIT:
                        formatting_retry_count += 1
                        logger.warning(
                            "document chunk formatting was not preserved; retrying chunk translation",
                            **failure_context,
                        )
                        continue
                    logger.warning(
                        "document chunk formatting was not preserved; accepting best-effort translation",
                        translated_payload_length=len(translated_payload),
                        **failure_context,
                    )
                    return [
                        _restore_translated_segment(
                            prepared_segment,
                            translated_segment,
                            strict=False,
                        )
                        for prepared_segment, translated_segment in zip(
                            prepared_segments, translated_segments, strict=True
                        )
                    ]
        except SegmentTranslationError as exc:
            failure_context = {
                "chunk_size": len(segments),
                "payload_length": len(payload),
                "failure_reason": str(exc),
                "target_language": target_language.value,
                **chunk_stats,
            }
            if "translated_payload" in locals():
                failure_context["translated_payload_length"] = len(translated_payload)
            if len(segments) == 1:
                logger.warning(
                    "document chunk translation failed with no smaller fallback available",
                    **failure_context,
                )
                raise

            midpoint = len(segments) // 2
            logger.warning(
                "document chunk translation failed; retrying smaller chunks",
                left_chunk_size=midpoint,
                right_chunk_size=len(segments) - midpoint,
                **failure_context,
            )
            return [
                *self._translate_chunk_with_fallback(segments[:midpoint], target_language),
                *self._translate_chunk_with_fallback(segments[midpoint:], target_language),
            ]

    def _convert_office_document(
        self,
        input_path: Path,
        target_format: str,
        working_dir: Path,
    ) -> Path:
        profile_dir = working_dir / "soffice-profile"
        profile_dir.mkdir(exist_ok=True)
        output_path = working_dir / f"{input_path.stem}.{target_format}"
        if output_path.exists():
            output_path.unlink()

        command = [
            *shlex.split(self._converter_command),
            "--headless",
            f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
            "--convert-to",
            target_format,
            "--outdir",
            str(working_dir),
            str(input_path),
        ]
        try:
            completed = subprocess.run(
                command,
                check=True,
                capture_output=True,
                text=True,
                timeout=self._conversion_timeout_seconds,
            )
        except FileNotFoundError:
            raise RuntimeError(
                f"Configured office converter command was not found: {self._converter_command}"
            ) from None
        except subprocess.CalledProcessError as exc:
            logger.error(
                "office conversion failed",
                command=command,
                stdout=exc.stdout[-4000:],
                stderr=exc.stderr[-4000:],
                returncode=exc.returncode,
            )
            raise RuntimeError(
                f"Failed to convert {input_path.suffix or 'office document'} to {target_format}"
            ) from None

        if output_path.exists():
            logger.info(
                "office conversion finished",
                input_file=str(input_path),
                output_file=str(output_path),
                stdout=completed.stdout[-2000:],
                stderr=completed.stderr[-2000:],
            )
            return output_path

        raise RuntimeError(
            f"Office converter completed but no {target_format} output was produced for {input_path.name}"
        )


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _iter_document_paragraphs(document: DocumentObject) -> Iterable[Paragraph]:
    seen: set[int] = set()

    def yield_paragraphs(paragraphs: Sequence[Paragraph]) -> Iterable[Paragraph]:
        for paragraph in paragraphs:
            paragraph_id = id(paragraph._element)
            if paragraph_id in seen:
                continue
            seen.add(paragraph_id)
            yield paragraph

    yield from yield_paragraphs(document.paragraphs)
    for table in document.tables:
        yield from _iter_table_paragraphs(table, seen)

    for section in document.sections:
        yield from yield_paragraphs(section.header.paragraphs)
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table, seen)
        yield from yield_paragraphs(section.footer.paragraphs)
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table, seen)


def _iter_table_paragraphs(table: Table, seen: set[int]) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph_id = id(paragraph._element)
                if paragraph_id in seen:
                    continue
                seen.add(paragraph_id)
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table, seen)


def _replace_paragraph_text(paragraph: Paragraph, translated_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = translated_text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(translated_text)


def _make_segment_marker() -> str:
    return f"[[PAPERLAAS_SEGMENT_BREAK_{uuid4().hex}]]"


def _prepare_segment_for_translation(segment: str) -> _PreparedSegment:
    match = _LEADING_TRAILING_WHITESPACE_RE.match(segment)
    if match is None:
        return _PreparedSegment("", segment, "", segment.count("\t"), segment.count("\n"))

    leading_whitespace, body, trailing_whitespace = match.groups()
    encoded_body = body.replace("\t", _TAB_TOKEN).replace("\n", _LINE_BREAK_TOKEN)
    return _PreparedSegment(
        leading_whitespace=leading_whitespace,
        encoded_body=encoded_body,
        trailing_whitespace=trailing_whitespace,
        tab_count=body.count("\t"),
        line_break_count=body.count("\n"),
    )


def _restore_translated_segment(
    prepared_segment: _PreparedSegment,
    translated_segment: str,
    *,
    strict: bool = True,
) -> str:
    if strict and translated_segment.count(_TAB_TOKEN) != prepared_segment.tab_count:
        raise _FormattingPreservationError(
            "Model did not preserve tab formatting tokens during translation"
        )
    if strict and translated_segment.count(_LINE_BREAK_TOKEN) != prepared_segment.line_break_count:
        raise _FormattingPreservationError(
            "Model did not preserve line break formatting tokens during translation"
        )

    restored_body = translated_segment.replace(_TAB_TOKEN, "\t").replace(_LINE_BREAK_TOKEN, "\n")
    return (
        prepared_segment.leading_whitespace
        + restored_body
        + prepared_segment.trailing_whitespace
    )


def _split_translated_payload(
    translated_payload: str,
    marker: str,
    expected_segments: int,
) -> list[str]:
    if expected_segments == 1:
        return [translated_payload]

    marker_pattern = re.compile(rf"\s*{re.escape(marker)}\s*")
    segments = marker_pattern.split(translated_payload)
    if len(segments) != expected_segments:
        exact_marker_count = translated_payload.count(marker)
        marker_like_count = len(_SEGMENT_MARKER_TOKEN_RE.findall(translated_payload))
        raise SegmentTranslationError(
            "Model did not preserve document segment boundaries during translation "
            f"(expected_segments={expected_segments}, expected_marker_count={expected_segments - 1}, "
            f"exact_marker_count={exact_marker_count}, marker_like_count={marker_like_count})"
        )
    return segments


def _get_chunk_stats(segments: Sequence[str]) -> dict[str, int]:
    lengths = [len(segment) for segment in segments]
    if not lengths:
        return {
            "min_segment_length": 0,
            "max_segment_length": 0,
            "avg_segment_length": 0,
        }
    return {
        "min_segment_length": min(lengths),
        "max_segment_length": max(lengths),
        "avg_segment_length": sum(lengths) // len(lengths),
    }
