from io import BytesIO
import re

from docx import Document
import pytest

from paperlaas_translate.llm import SegmentTranslationError
from paperlaas_translate.models import TargetLanguage
from paperlaas_translate.translators.documents import (
    OfficeAndTextTranslator,
    _prepare_segment_for_translation,
    _restore_translated_segment,
    _split_translated_payload,
)


class FakeSegmentTranslator:
    def __init__(self) -> None:
        self.calls: list[list[str]] = []

    def translate_segments(self, segments: list[str], target_language: TargetLanguage) -> list[str]:
        self.calls.append(segments)
        translated_segments: list[str] = []
        marker_re = re.compile(r"(\[\[PAPERLAAS_[^\]]+\]\])")
        for segment in segments:
            parts = marker_re.split(segment)
            translated_segments.append(
                "".join(
                    part if marker_re.fullmatch(part) else f"{target_language.iso_code}:{part}"
                    for part in parts
                )
            )
        return translated_segments


def test_docx_translation_preserves_structure() -> None:
    segment_translator = FakeSegmentTranslator()
    document = Document()
    document.add_paragraph("Hello world")
    table = document.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Cell content"
    buffer = BytesIO()
    document.save(buffer)

    translator = OfficeAndTextTranslator(segment_translator, batch_chars=1000, batch_items=10)
    translated_bytes = translator.translate(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
        TargetLanguage.GERMAN,
    )

    translated_document = Document(BytesIO(translated_bytes))
    assert translated_document.paragraphs[0].text == "de:Hello world"
    assert translated_document.tables[0].rows[0].cells[0].text == "de:Cell content"
    assert len(segment_translator.calls) == 1
    assert len(segment_translator.calls[0]) == 1


def test_text_translation_preserves_blank_lines() -> None:
    translator = OfficeAndTextTranslator(FakeSegmentTranslator(), batch_chars=1000, batch_items=10)

    translated = translator.translate(
        "text/plain",
        b"First paragraph\n\nSecond paragraph",
        TargetLanguage.FRENCH,
    )

    assert translated.decode("utf-8") == "fr:First paragraph\n\nfr:Second paragraph"


def test_docx_translation_preserves_tabs_in_paragraph_text() -> None:
    translator = OfficeAndTextTranslator(FakeSegmentTranslator(), batch_chars=1000, batch_items=10)

    document = Document()
    document.add_paragraph("\t\t\t\tDate: 2026-03-21")
    buffer = BytesIO()
    document.save(buffer)

    translated_bytes = translator.translate(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
        TargetLanguage.GERMAN,
    )

    translated_document = Document(BytesIO(translated_bytes))
    assert translated_document.paragraphs[0].text == "\t\t\t\tde:Date: 2026-03-21"


class FlakyBatchSegmentTranslator:
    def __init__(self) -> None:
        self.calls: list[list[str]] = []

    def translate_segments(self, segments: list[str], target_language: TargetLanguage) -> list[str]:
        self.calls.append(segments)
        if re.search(r"\[\[PAPERLAAS_SEGMENT_BREAK_[0-9a-f]+\]\]", segments[0]):
            raise SegmentTranslationError(
                "Model did not preserve document segment boundaries during translation"
            )
        return [f"{target_language.iso_code}:{segments[0]}"]


def test_text_translation_splits_batches_after_segment_count_mismatch() -> None:
    segment_translator = FlakyBatchSegmentTranslator()
    translator = OfficeAndTextTranslator(
        segment_translator,
        batch_chars=1000,
        batch_items=10,
    )

    translated = translator.translate(
        "text/plain",
        b"First paragraph\n\nSecond paragraph\n\nThird paragraph",
        TargetLanguage.GERMAN,
    )

    assert translated.decode("utf-8") == (
        "de:First paragraph\n\nde:Second paragraph\n\nde:Third paragraph"
    )
    assert len(segment_translator.calls) >= 3


class FormattingRetrySegmentTranslator:
    def __init__(self, responses: list[str]) -> None:
        self.calls: list[list[str]] = []
        self._responses = responses

    def translate_segments(self, segments: list[str], target_language: TargetLanguage) -> list[str]:
        self.calls.append(segments)
        response_index = min(len(self.calls) - 1, len(self._responses) - 1)
        return [self._responses[response_index]]


def test_docx_translation_retries_when_formatting_tokens_are_not_preserved() -> None:
    segment_translator = FormattingRetrySegmentTranslator(
        [
            "de:Header Value",
            "de:Header[[PAPERLAAS_TAB]]de:Value",
        ]
    )
    translator = OfficeAndTextTranslator(
        segment_translator,
        batch_chars=1000,
        batch_items=10,
    )

    document = Document()
    document.add_paragraph("Header\tValue")
    buffer = BytesIO()
    document.save(buffer)

    translated_bytes = translator.translate(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
        TargetLanguage.GERMAN,
    )

    translated_document = Document(BytesIO(translated_bytes))
    assert translated_document.paragraphs[0].text == "de:Header\tde:Value"
    assert len(segment_translator.calls) == 2


def test_docx_translation_accepts_best_effort_result_after_formatting_retries() -> None:
    segment_translator = FormattingRetrySegmentTranslator(
        [
            "de:Header Value",
            "de:Header Value",
            "de:Header Value",
        ]
    )
    translator = OfficeAndTextTranslator(
        segment_translator,
        batch_chars=1000,
        batch_items=10,
    )

    document = Document()
    document.add_paragraph("Header\tValue")
    buffer = BytesIO()
    document.save(buffer)

    translated_bytes = translator.translate(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
        TargetLanguage.GERMAN,
    )

    translated_document = Document(BytesIO(translated_bytes))
    assert translated_document.paragraphs[0].text == "de:Header Value"
    assert len(segment_translator.calls) == 3


def test_odt_translation_converts_via_docx(monkeypatch: pytest.MonkeyPatch, tmp_path) -> None:
    source_document = Document()
    source_document.add_paragraph("Hello world")
    source_document.add_paragraph("Second paragraph")
    source_buffer = BytesIO()
    source_document.save(source_buffer)

    conversion_calls: list[tuple[str, str]] = []

    def fake_convert(self, input_path, target_format, working_dir):
        conversion_calls.append((input_path.suffix, target_format))
        if input_path.suffix == ".odt" and target_format == "docx":
            converted_docx_path = working_dir / "source.docx"
            converted_docx_path.write_bytes(source_buffer.getvalue())
            return converted_docx_path
        if input_path.suffix == ".docx" and target_format == "odt":
            converted_odt_path = working_dir / "translated.odt"
            converted_odt_path.write_bytes(input_path.read_bytes())
            return converted_odt_path
        raise AssertionError(f"Unexpected conversion request: {input_path} -> {target_format}")

    monkeypatch.setattr(OfficeAndTextTranslator, "_convert_office_document", fake_convert)

    translator = OfficeAndTextTranslator(
        FakeSegmentTranslator(),
        batch_chars=1000,
        batch_items=10,
        temp_dir=str(tmp_path),
    )

    translated_bytes = translator.translate(
        "application/vnd.oasis.opendocument.text",
        b"fake odt payload",
        TargetLanguage.GERMAN,
    )

    translated_document = Document(BytesIO(translated_bytes))
    assert translated_document.paragraphs[0].text == "de:Hello world"
    assert translated_document.paragraphs[1].text == "de:Second paragraph"
    assert conversion_calls == [(".odt", "docx"), (".docx", "odt")]


def test_split_translated_payload_error_reports_marker_counts() -> None:
    marker = "[[PAPERLAAS_SEGMENT_BREAK_test]]"

    with pytest.raises(SegmentTranslationError) as exc_info:
        _split_translated_payload("translated without markers", marker, 3)

    assert "expected_segments=3" in str(exc_info.value)
    assert "exact_marker_count=0" in str(exc_info.value)
    assert "marker_like_count=0" in str(exc_info.value)


def test_prepare_and_restore_segment_preserves_whitespace_formatting() -> None:
    prepared_segment = _prepare_segment_for_translation("\t\tHeader\tValue\nLine 2  ")

    assert prepared_segment.leading_whitespace == "\t\t"
    assert prepared_segment.trailing_whitespace == "  "

    restored = _restore_translated_segment(
        prepared_segment,
        "de:Header[[PAPERLAAS_TAB]]de:Value[[PAPERLAAS_LINE_BREAK]]de:Line 2",
    )

    assert restored == "\t\tde:Header\tde:Value\nde:Line 2  "
